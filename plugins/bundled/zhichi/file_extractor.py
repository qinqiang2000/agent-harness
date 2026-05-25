"""智齿附件文字提取 - 支持 PDF 和 DOCX。"""

import logging
import tempfile
from pathlib import Path
from urllib.parse import urlparse

import httpx

logger = logging.getLogger(__name__)

_TIMEOUT = 30.0
_MAX_BYTES = 20 * 1024 * 1024  # 20MB
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx", ".doc"}
_MAX_PDF_CHARS = 8000  # 避免超长文本撑爆 context


class FileExtractError(Exception):
    pass


def _url_ext(url: str) -> str:
    path = urlparse(url).path
    return Path(path).suffix.lower()


def is_pdf_url(url: str) -> bool:
    return _url_ext(url) in _PDF_EXTS


def is_docx_url(url: str) -> bool:
    return _url_ext(url) in _DOCX_EXTS


async def _download(url: str) -> bytes:
    async with httpx.AsyncClient(timeout=_TIMEOUT, follow_redirects=True) as client:
        resp = await client.get(url)
        resp.raise_for_status()
    content = resp.content
    if len(content) > _MAX_BYTES:
        raise FileExtractError(f"文件超过 20MB 上限: {url}")
    return content


def _extract_pdf(data: bytes) -> str:
    try:
        import pymupdf
    except ImportError:
        raise FileExtractError("pymupdf 未安装，无法解析 PDF")

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as f:
        f.write(data)
        f.flush()
        doc = pymupdf.open(f.name)
        parts = []
        for page in doc:
            text = page.get_text().strip()
            if text:
                parts.append(text)
        doc.close()

    full = "\n\n".join(parts).strip()
    if not full:
        raise FileExtractError("PDF 无可提取文字（可能是扫描件图片）")
    if len(full) > _MAX_PDF_CHARS:
        full = full[:_MAX_PDF_CHARS] + "\n…（内容过长，已截断）"
    return full


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(data))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
    except ImportError:
        raise FileExtractError("python-docx 未安装，无法解析 DOCX")
    except Exception as e:
        raise FileExtractError(f"DOCX 解析失败: {e}")

    full = "\n".join(parts).strip()
    if not full:
        raise FileExtractError("DOCX 无可提取文字")
    if len(full) > _MAX_PDF_CHARS:
        full = full[:_MAX_PDF_CHARS] + "\n…（内容过长，已截断）"
    return full


async def extract_file_text(url: str) -> str:
    """下载并提取文件文字，返回纯文本。失败抛 FileExtractError。"""
    ext = _url_ext(url)
    logger.info(f"[FileExtractor] Downloading {ext} file: {url}")
    data = await _download(url)

    if ext in _PDF_EXTS:
        text = _extract_pdf(data)
    elif ext in _DOCX_EXTS:
        text = _extract_docx(data)
    else:
        raise FileExtractError(f"不支持的文件类型: {ext}")

    logger.info(f"[FileExtractor] Extracted {len(text)} chars from {ext}")
    return text
